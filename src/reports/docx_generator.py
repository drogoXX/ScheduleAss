"""
DOCX Executive Summary Report Generator
Generates professional Word documents with schedule analysis

Layout, typography and colour come from ``src/reports/docx_styles.py`` — the
shared design system also used by ``tools/build_user_guide_docx.py`` — and the
charts come from ``src/reports/docx_charts.py``, which renders them through the
same Plotly theme as the dashboard. Nothing here defines its own colours or
fonts, so the report matches the screen by construction.
"""

from datetime import datetime
from typing import Dict, List
import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from src.reports import docx_charts
from src.reports.docx_styles import (
    ACCENT_FILL,
    CRITICAL_FILL,
    GRAY_MEDIUM,
    SUCCESS_FILL,
    WARNING_FILL,
    add_body,
    add_bullets,
    add_cover_page,
    add_figure,
    add_note,
    add_table,
    add_toc,
    configure_headers_footers,
    configure_styles,
    rating_rgb,
    status_fill,
)
from src.ui.palette import (
    fmt_count,
    fmt_days,
    fmt_index,
    fmt_pct,
    fmt_score,
)

# Status glyph + the fill that goes behind it, so pass/fail reads at a glance
# instead of being black text in a black-text table.
_STATUS_MARKS = {
    'pass': ('✓ PASS', SUCCESS_FILL),
    'fail': ('✗ FAIL', CRITICAL_FILL),
    'n/a': ('○ N/A', None),
    'manual': ('◐ MANUAL', WARNING_FILL),
}


class DOCXGenerator:
    """Generates executive summary reports in DOCX format"""

    def __init__(self, project_name: str, schedule_data: Dict, analysis_results: Dict):
        """
        Initialize DOCX generator

        Args:
            project_name: Name of the project
            schedule_data: Parsed schedule data
            analysis_results: Complete analysis results including metrics and recommendations
        """
        self.project_name = project_name
        self.schedule_data = schedule_data
        self.analysis_results = analysis_results
        self.document = Document()
        self._generated_at = datetime.now()
        self._figure_number = 0

    # ------------------------------------------------------------------
    # Entry point
    # ------------------------------------------------------------------
    def generate(self) -> bytes:
        """Generate the DOCX report and return as bytes"""
        # Set up document properties
        self._setup_document()

        # Add content sections
        self._add_cover_page()
        self._add_table_of_contents()
        self._add_executive_summary()
        self._add_dcma_compliance()
        self._add_missing_logic_breakdown()
        self._add_key_metrics()
        self._add_wbs_analysis()
        self._add_issues_summary()
        self._add_recommendations()
        self._add_appendix()

        # Running header and footer are applied last so they cover every
        # section created above.
        self._configure_headers_footers()

        # Save to bytes
        file_stream = io.BytesIO()
        self.document.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------
    @property
    def _doc_number(self) -> str:
        return f"SQA-{self._generated_at.strftime('%Y%m%d')}"

    @property
    def _doc_date(self) -> str:
        return self._generated_at.strftime('%d %B %Y')

    def _add_figure(self, image_stream, caption: str) -> bool:
        """
        Place a chart with an auto-numbered caption.

        Returns False when the image could not be rendered, so the caller can
        skip any surrounding prose that assumes a figure is present.
        """
        if image_stream is None:
            return False
        self._figure_number += 1
        add_figure(self.document, image_stream,
                   caption=f"Figure {self._figure_number} — {caption}")
        return True

    def _setup_document(self):
        """Set up document styles and properties"""
        configure_styles(self.document)

        # Set document title
        self.document.core_properties.title = f"Schedule Quality Analysis - {self.project_name}"
        self.document.core_properties.author = "Schedule Quality Analyzer"
        self.document.core_properties.created = self._generated_at

    def _configure_headers_footers(self):
        configure_headers_footers(
            self.document,
            project=self.project_name,
            doc_number=self._doc_number,
            doc_date=self._doc_date,
        )

    # ------------------------------------------------------------------
    # Sections
    # ------------------------------------------------------------------
    def _add_cover_page(self):
        """Add cover page"""
        health_score = self.analysis_results['performance_metrics']['health_score']
        stats = self.analysis_results['performance_metrics'].get('statistics', {})
        rating = health_score.get('rating', 'Unknown')

        add_cover_page(
            self.document,
            title="Schedule Quality Analysis",
            subtitle="DCMA 14-Point Schedule Assessment",
            meta=[
                ("Project", self.project_name),
                ("Document No.", self._doc_number),
                ("Date", self._doc_date),
                ("Activities", fmt_count(stats.get('total_activities', 0))),
                ("Prepared by", "Schedule Quality Analyzer"),
                ("Classification", "Internal Use"),
            ],
            headline=(
                fmt_score(health_score.get('score', 0)),
                f"Schedule Health Score — {rating}",
                rating_rgb(rating),
            ),
        )

    def _add_table_of_contents(self):
        add_toc(self.document)

    def _add_executive_summary(self):
        """Add executive summary section"""
        self.document.add_heading('Executive Summary', level=1)

        # Overall assessment
        health_score = self.analysis_results['performance_metrics']['health_score']
        stats = self.analysis_results['performance_metrics']['statistics']

        add_body(
            self.document,
            f"This report presents an analysis of the {self.project_name} schedule "
            f"against the industry-standard DCMA 14-Point Schedule Assessment. "
            f"The schedule contains {fmt_count(stats['total_activities'])} activities "
            f"with an overall health score of {fmt_score(health_score['score'])}, "
            f"rated as \"{health_score['rating']}\"."
        )

        # Any ceiling that limited the score must be stated, otherwise the
        # number looks arbitrary.
        for cap in health_score.get('caps', []) or []:
            add_note(self.document, "Score limited", str(cap), fill=WARNING_FILL)

        # Key findings
        self.document.add_heading('Key Findings', level=2)

        issues = self.analysis_results.get('issues', [])
        high_priority_issues = [i for i in issues if i.get('severity') == 'high']

        if high_priority_issues:
            add_body(self.document,
                     f"{fmt_count(len(high_priority_issues))} high-priority "
                     f"issue(s) require attention:")
            add_bullets(self.document,
                        [issue.get('title', '') for issue in high_priority_issues[:5]])
            if len(high_priority_issues) > 5:
                add_body(
                    self.document,
                    f"A further {fmt_count(len(high_priority_issues) - 5)} "
                    f"high-priority issue(s) are listed in the Issues Summary.",
                )
        else:
            add_bullets(self.document, ['No high-priority issues identified.'])

        self._add_figure(
            docx_charts.issues_by_severity_chart(issues),
            "Issues by severity",
        )

    def _add_dcma_compliance(self):
        """Add DCMA 14-Point compliance checklist with all categories"""
        self.document.add_heading('DCMA 14-Point Compliance Assessment', level=1)

        # Get DCMA 14-point summary
        dcma_14 = self.analysis_results.get('dcma_14_point', {})

        if not dcma_14:
            # Fallback to old format if dcma_14_point not available
            self._add_dcma_compliance_legacy()
            return

        # Overall Score
        add_body(
            self.document,
            f"{fmt_count(dcma_14.get('overall_pass_count', 0))} PASS   |   "
            f"{fmt_count(dcma_14.get('overall_fail_count', 0))} FAIL   |   "
            f"{fmt_count(dcma_14.get('overall_na_count', 0))} N/A   |   "
            f"{fmt_count(dcma_14.get('overall_manual_count', 0))} MANUAL",
            bold_lead=f"Overall Score: {dcma_14.get('overall_score_text', 'N/A')}   —   ",
        )

        self._add_figure(
            docx_charts.dcma_compliance_chart(dcma_14),
            "DCMA 14-point assessment outcome",
        )

        # Iterate through categories
        categories = dcma_14.get('categories', {})

        for cat_name, cat_data in categories.items():
            # Category header
            self.document.add_heading(
                f"Category {cat_data['number']}: {cat_name}", level=2)

            # Calculate category pass/fail
            cat_metrics = cat_data['metrics']
            cat_pass = sum(1 for m in cat_metrics if m['status'] == 'pass')
            cat_total = len([m for m in cat_metrics
                             if m['status'] not in ['n/a', 'manual', 'unknown']])

            cat_score_text = f"[{cat_pass}/{cat_total}]" if cat_total > 0 else "[N/A]"
            add_body(self.document, "", bold_lead=f"Category Score: {cat_score_text}")

            rows = []
            for metric in cat_metrics:
                mark, _ = _STATUS_MARKS.get(
                    str(metric['status']).lower(), ('? UNKNOWN', None))
                rows.append([
                    str(metric['number']),
                    metric['name'],
                    mark,
                    metric['description'],
                ])

            add_table(
                self.document,
                ['#', 'Metric', 'Status', 'Result'],
                rows,
                widths=[Cm(1.2), Cm(4.6), Cm(2.4), Cm(8.3)],
                # Tint by outcome: the status column is the point of the table.
                row_fill=lambda row: status_fill(
                    'pass' if 'PASS' in row[2]
                    else 'fail' if 'FAIL' in row[2]
                    else 'warning' if 'MANUAL' in row[2]
                    else 'n/a'
                ),
            )

            # Add recommendations for failed metrics
            failed_metrics = [m for m in cat_metrics
                              if m['status'] == 'fail' and m.get('recommendation')]
            if failed_metrics:
                self.document.add_heading('Recommendations', level=3)
                add_bullets(
                    self.document,
                    [(f"{m['name']}: ", m['recommendation']) for m in failed_metrics],
                )

        # Add legend
        self.document.add_heading('Status Legend', level=2)
        add_bullets(self.document, [
            ('✓ PASS', ' — Meets DCMA standard'),
            ('✗ FAIL', ' — Does not meet DCMA standard (action required)'),
            ('○ N/A', ' — Not applicable or data not available'),
            ('◐ MANUAL', ' — Manual verification required'),
        ])

    def _add_dcma_compliance_legacy(self):
        """Legacy DCMA compliance section (fallback)"""
        dcma_metrics = self.analysis_results['dcma_metrics']

        checklist = [
            ('Negative Lags', dcma_metrics.get('negative_lags', {}).get('status', 'unknown'),
             f"{fmt_count(dcma_metrics.get('negative_lags', {}).get('count', 0))} found (Target: 0)"),
            ('Positive Lags', dcma_metrics.get('positive_lags', {}).get('status', 'unknown'),
             f"{fmt_pct(dcma_metrics.get('positive_lags', {}).get('percentage', 0))} (Target: ≤5%)"),
            ('Hard Constraints', dcma_metrics.get('hard_constraints', {}).get('status', 'unknown'),
             f"{fmt_pct(dcma_metrics.get('hard_constraints', {}).get('percentage', 0))} (Target: ≤10%)"),
            ('Missing Logic', dcma_metrics.get('missing_logic', {}).get('status', 'unknown'),
             f"{fmt_count(dcma_metrics.get('missing_logic', {}).get('count', 0))} activities (Target: 0)"),
        ]

        add_table(
            self.document,
            ['Metric', 'Status', 'Result'],
            [[name, status.upper(), result] for name, status, result in checklist],
            widths=[Cm(5.0), Cm(3.0), Cm(8.5)],
            row_fill=lambda row: status_fill(row[1].lower()),
        )

    def _add_missing_logic_breakdown(self):
        """Add detailed missing logic breakdown section"""
        self.document.add_heading('Logic Completeness Analysis', level=1)

        dcma_metrics = self.analysis_results.get('dcma_metrics', {})
        missing_logic_info = dcma_metrics.get('missing_logic', {})

        # Check if there's any missing logic
        total_missing = missing_logic_info.get('count', 0)

        if total_missing == 0:
            add_note(
                self.document, "Complete",
                "All activities have complete logic relationships "
                "(predecessors and successors).",
                fill=SUCCESS_FILL,
            )
            return

        # Add overview paragraph
        add_body(
            self.document,
            "Activities without proper predecessors or successors indicate "
            "incomplete schedule logic and can affect critical path calculations.",
            bold_lead=f"Found {fmt_count(total_missing)} activities with missing "
                      f"logic relationships. ",
        )

        # Create breakdown table
        self.document.add_heading('Missing Logic Breakdown', level=2)

        pred_only = missing_logic_info.get('missing_predecessor_only_count', 0)
        succ_only = missing_logic_info.get('missing_successor_only_count', 0)
        both_count = missing_logic_info.get('missing_both_count', 0)

        add_table(
            self.document,
            ['Category', 'Count', 'Description'],
            [
                ['Total Missing Logic', fmt_count(total_missing),
                 'Unique activities with missing predecessor and/or successor'],
                ['  ├─ Missing Predecessor Only', fmt_count(pred_only),
                 'Activities that need predecessors added'],
                ['  ├─ Missing Successor Only', fmt_count(succ_only),
                 'Activities that need successors added'],
                ['  └─ Missing Both', fmt_count(both_count),
                 'Activities that need both predecessors and successors'],
            ],
            widths=[Cm(5.6), Cm(2.0), Cm(8.9)],
        )

        # DCMA validation section
        self.document.add_heading('DCMA Validation', level=2)

        dcma_pred = dcma_metrics.get('dcma_missing_predecessors', {}).get('count', 0)
        dcma_succ = dcma_metrics.get('dcma_missing_successors', {}).get('count', 0)

        add_bullets(self.document, [
            ('DCMA Missing Predecessors: ',
             f"{fmt_count(dcma_pred)} activities "
             f"(includes {fmt_count(pred_only)} pred-only + "
             f"{fmt_count(both_count)} both)"),
            ('DCMA Missing Successors: ',
             f"{fmt_count(dcma_succ)} activities "
             f"(includes {fmt_count(succ_only)} succ-only + "
             f"{fmt_count(both_count)} both)"),
        ])

        add_note(
            self.document, "Note",
            "Activities missing both predecessors and successors are counted in "
            "each DCMA category. Breakdown validation: "
            f"{pred_only} + {succ_only} + {both_count} = {total_missing} total.",
        )

        # Add recommendation if there are issues
        add_note(
            self.document, "Action required",
            "Add logical relationships to all activities. Every activity "
            "(except start/finish milestones) should have both predecessors and "
            "successors to ensure proper schedule logic and accurate critical "
            "path calculations.",
            fill=WARNING_FILL,
        )

    def _add_key_metrics(self):
        """Add key performance metrics"""
        self.document.add_heading('Key Performance Metrics', level=1)

        perf_metrics = self.analysis_results['performance_metrics']
        cpli = perf_metrics.get('cpli', {})
        bei = perf_metrics.get('bei', {})

        add_table(
            self.document,
            ['Metric', 'Value', 'Target', 'Status'],
            [
                ['Critical Path Length Index (CPLI)',
                 fmt_index(cpli.get('value', 0)), '≥ 0.950',
                 str(cpli.get('status', 'unknown')).upper()],
                ['Baseline Execution Index (BEI)',
                 fmt_index(bei.get('value', 0)), '≥ 0.950',
                 str(bei.get('status', 'unknown')).upper()],
            ],
            widths=[Cm(7.5), Cm(3.0), Cm(3.0), Cm(3.0)],
            row_fill=lambda row: status_fill(row[3].lower()),
        )

        if cpli.get('description'):
            add_body(self.document, str(cpli['description']),
                     bold_lead="CPLI: ")
        add_body(
            self.document,
            f"{fmt_count(bei.get('completed', 0))} completed of "
            f"{fmt_count(bei.get('planned', 0))} planned.",
            bold_lead="BEI: ",
        )

    def _add_wbs_analysis(self):
        """Add WBS (Work Breakdown Structure) analysis"""
        # Check if WBS analysis is available
        wbs_analysis = self.analysis_results['dcma_metrics'].get('wbs_analysis', {})

        if not wbs_analysis.get('available'):
            # Skip this section if WBS data is not available
            return

        self.document.add_heading('WBS Analysis', level=1)

        # Overview
        add_table(
            self.document,
            ['Measure', 'Value'],
            [
                ['Total Activities', fmt_count(wbs_analysis.get('total_activities', 0))],
                ['Activities with WBS Codes',
                 fmt_count(wbs_analysis.get('activities_with_wbs', 0))],
                ['Average WBS Depth',
                 fmt_days(wbs_analysis.get('avg_depth', 0), unit=False)],
                ['Maximum WBS Depth', fmt_count(wbs_analysis.get('max_depth', 0))],
            ],
            widths=[Cm(8.5), Cm(8.0)],
        )

        # WBS Level 1 (Phases) Analysis
        level1 = wbs_analysis.get('level_1_phases', {})
        if level1:
            self.document.add_heading('WBS Level 1 — Phases', level=2)

            shown, total = docx_charts.phase_count_used(level1)
            caption = "Health score by WBS phase"
            if total > shown:
                caption += (f" (the {shown} phases with the most activities, "
                            f"of {total})")
            self._add_figure(docx_charts.wbs_health_chart(level1), caption)

            rows = []
            for wbs_code, stats in sorted(level1.items()):
                health_score = stats.get('health_score', {})
                rows.append([
                    f"Phase {wbs_code}",
                    fmt_count(stats.get('activity_count', 0)),
                    fmt_days(stats.get('avg_float', 0), unit=False),
                    fmt_count(stats.get('critical_count', 0)),
                    fmt_count(stats.get('negative_float_count', 0)),
                    fmt_score(health_score.get('score', 0)),
                    health_score.get('rating', 'Unknown'),
                ])

            add_table(
                self.document,
                ['Phase', 'Activities', 'Avg Float', 'Critical', 'Negative',
                 'Health Score', 'Rating'],
                rows,
                highlight=lambda row: row[6] in ('Poor', 'Critical'),
            )

        # WBS Level 2 (Areas) Analysis
        level2 = wbs_analysis.get('level_2_areas', {})
        if level2:
            self.document.add_heading('WBS Level 2 — Areas', level=2)

            # Sorted by health score so problem areas come first
            areas_with_scores = []
            for wbs_code, stats in level2.items():
                health_score = stats.get('health_score', {})
                areas_with_scores.append(
                    (wbs_code, stats, health_score.get('score', 0)))
            areas_with_scores.sort(key=lambda x: x[2])

            rows = []
            for wbs_code, stats, _ in areas_with_scores:
                health_score = stats.get('health_score', {})
                activity_count = stats.get('activity_count', 0)
                critical_count = stats.get('critical_count', 0)
                pct_critical = ((critical_count / activity_count * 100)
                                if activity_count > 0 else 0)

                rows.append([
                    f"Area {wbs_code}",
                    fmt_count(activity_count),
                    fmt_days(stats.get('avg_float', 0), unit=False),
                    fmt_count(critical_count),
                    fmt_pct(pct_critical),
                    fmt_score(health_score.get('score', 0)),
                    health_score.get('rating', 'Unknown'),
                ])

            add_table(
                self.document,
                ['Area', 'Activities', 'Avg Float', 'Critical', '% Critical',
                 'Health Score', 'Rating'],
                rows,
                highlight=lambda row: row[6] in ('Poor', 'Critical'),
            )

        # WBS Health Score Legend
        self.document.add_heading('Health Score Interpretation', level=2)
        add_body(self.document, "Health scores are calculated from:")
        add_bullets(self.document, [
            ('Critical Path % (40 points): ', 'lower is better'),
            ('Average Float (30 points): ', 'higher is better'),
            ('Negative Float % (20 points): ', 'lower is better'),
            ('Activity Distribution (10 points): ', 'balanced is better'),
        ])

        add_table(
            self.document,
            ['Rating', 'Score', 'Interpretation'],
            [
                ['Excellent', '80-100', 'Well-balanced, low risk'],
                ['Good', '65-79', 'Acceptable performance'],
                ['Fair', '50-64', 'Some concerns'],
                ['Poor', '35-49', 'Significant issues'],
                ['Critical', '0-34', 'Immediate attention required'],
            ],
            widths=[Cm(3.5), Cm(3.0), Cm(10.0)],
        )

    def _add_issues_summary(self):
        """Add issues summary table"""
        self.document.add_heading('Issues Summary', level=1)

        issues = self.analysis_results.get('issues', [])

        if not issues:
            add_note(self.document, "Clear", "No issues identified.",
                     fill=SUCCESS_FILL)
            return

        severity_order = {'high': 0, 'medium': 1, 'low': 2}
        ordered = sorted(
            issues,
            key=lambda i: severity_order.get(str(i.get('severity', '')).lower(), 3),
        )

        add_table(
            self.document,
            ['Priority', 'Category', 'Issue', 'Count'],
            [[
                str(issue.get('severity', '')).upper(),
                issue.get('category', ''),
                issue.get('title', ''),
                fmt_count(issue.get('count', 0)),
            ] for issue in ordered],
            widths=[Cm(2.2), Cm(3.4), Cm(8.4), Cm(2.0)],
            highlight=lambda row: row[0] == 'HIGH',
        )

    def _add_recommendations(self):
        """Add recommendations section"""
        self.document.add_heading('Recommendations', level=1)

        recommendations = self.analysis_results.get('recommendations', [])

        if not recommendations:
            add_body(self.document, 'No recommendations at this time.')
            return

        # Group by priority
        high = [r for r in recommendations if r.get('priority') == 'high']
        medium = [r for r in recommendations if r.get('priority') == 'medium']

        # High priority
        if high:
            self.document.add_heading('High Priority', level=2)
            for i, rec in enumerate(high, 1):
                self.document.add_heading(f"{i}. {rec.get('title', '')}", level=3)
                add_body(self.document, rec.get('description', ''),
                         bold_lead="Description: ")
                add_body(self.document, rec.get('recommendation', ''),
                         bold_lead="Recommendation: ")
                add_body(self.document,
                         f"{rec.get('impact', '')}   |   Effort: {rec.get('effort', '')}",
                         bold_lead="Impact: ")

        # Medium priority
        if medium:
            self.document.add_heading('Medium Priority', level=2)
            for i, rec in enumerate(medium, 1):
                self.document.add_heading(f"{i}. {rec.get('title', '')}", level=3)
                add_body(self.document, rec.get('recommendation', ''),
                         bold_lead="Recommendation: ")

    def _add_appendix(self):
        """Add appendix with methodology"""
        self.document.add_page_break()
        self.document.add_heading('Appendix: Methodology', level=1)

        self.document.add_heading('Framework', level=2)
        add_body(
            self.document,
            "Defense Contract Management Agency best practices for schedule "
            "quality, focusing on logic integrity, schedule realism and "
            "execution readiness.",
            bold_lead="DCMA 14-Point Schedule Assessment. ",
        )

        self.document.add_heading('Schedule Health Score', level=2)
        add_body(
            self.document,
            "The 0-100 score is a weighted average of the DCMA checks above. "
            "Each check scores 100 at or better than its DCMA target and "
            "declines linearly to 0 at a defined bound. Checks without "
            "supporting data are excluded and the remaining weights "
            "renormalised."
        )
        add_note(
            self.document, "Scope",
            "Thresholds follow the DCMA 14-Point Assessment. The weights are "
            "this application's assessment of relative severity and are not "
            "defined by DCMA.",
        )

        self.document.add_heading('Metric Definitions', level=2)
        add_table(
            self.document,
            ['Metric', 'Formula', 'Target'],
            [
                ['CPLI — Critical Path Length Index',
                 '(Critical Path + Total Float) / Critical Path', '≥ 0.950'],
                ['BEI — Baseline Execution Index',
                 'Completed Tasks / Planned Tasks', '≥ 0.950'],
            ],
            widths=[Cm(6.0), Cm(7.5), Cm(3.0)],
        )
