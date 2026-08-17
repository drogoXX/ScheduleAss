"""
Shared Word design system.

Extracted from ``tools/build_user_guide_docx.py``, which produced the shipped
User Guide, so both that tool and ``src/reports/docx_generator.py`` render from
one definition and cannot drift apart.

Colours come from :mod:`src.ui.palette` — the same values the Streamlit UI and
the Plotly charts use — converted here to the ``RGBColor`` and bare-hex forms
python-docx and OOXML need.
"""

from typing import Iterable, Optional, Sequence

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.ui.palette import BODY_FONT_NAME, COLORS, RATING_COLORS

BODY_FONT = BODY_FONT_NAME


def _hex(value: str) -> str:
    """'#004B87' -> '004B87' (the form OOXML fill attributes want)."""
    return value.lstrip("#").upper()


def _rgb(value: str) -> RGBColor:
    """'#004B87' -> RGBColor(0x00, 0x4B, 0x87)."""
    return RGBColor.from_string(_hex(value))


# ---------------------------------------------------------------------------
# Palette, in python-docx form
# ---------------------------------------------------------------------------
DOCX_COLORS = {name: _rgb(value) for name, value in COLORS.items()}

# Convenience aliases matching the roles used throughout the document.
PRIMARY = DOCX_COLORS["primary"]
ACCENT = DOCX_COLORS["accent"]
BLACK = DOCX_COLORS["text"]
GRAY_DARK = _rgb("#424242")
GRAY_MEDIUM = DOCX_COLORS["text_muted"]
WHITE = DOCX_COLORS["white"]

# Cell fills (bare hex, no '#').
HEADER_FILL = _hex(COLORS["primary"])
ALT_ROW_FILL = _hex(COLORS["surface"])
NOTE_FILL = _hex(COLORS["primary_light"])
ACCENT_FILL = _hex(COLORS["accent"])
CRITICAL_FILL = "FADBD8"   # Light red tint; readable behind black body text
WARNING_FILL = "FDEBD0"    # Light amber tint
SUCCESS_FILL = "E3F1E4"    # Light green tint
META_FILL = "D6E4F0"

PAGE_WIDTH_A4 = Cm(21.0)
PAGE_HEIGHT_A4 = Cm(29.7)
MARGIN_TOP = Cm(2.5)
MARGIN_BOTTOM = Cm(2.0)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(2.0)
CONTENT_WIDTH = Cm(16.5)


def rating_rgb(rating: str) -> RGBColor:
    """Font colour for a health-score rating, from the shared palette."""
    return _rgb(RATING_COLORS.get(str(rating).title(), COLORS["text_muted"]))


def status_fill(status: str) -> Optional[str]:
    """Cell fill for a pass/warning/fail token, or None to leave unshaded."""
    return {
        "pass": SUCCESS_FILL, "good": SUCCESS_FILL,
        "warning": WARNING_FILL,
        "fail": CRITICAL_FILL, "bad": CRITICAL_FILL,
    }.get(str(status).lower())


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def style_cell(cell, bold=False, size=Pt(9), color=None, align=None):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            run.font.name = BODY_FONT
            run.font.size = size
            run.font.bold = bold
            run.font.color.rgb = color or GRAY_DARK


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.makeelement(
        qn("w:shd"), {qn("w:fill"): hex_color, qn("w:val"): "clear"}
    )
    tc_pr.append(shading)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(borders.makeelement(
            qn(f"w:{name}"),
            {qn("w:val"): "none", qn("w:sz"): "0",
             qn("w:space"): "0", qn("w:color"): "auto"},
        ))
    tbl_pr.append(borders)


def add_field(paragraph, instruction):
    """Insert a Word field code (used for TOC and page numbers)."""
    run = paragraph.add_run()
    run._r.append(run._r.makeelement(
        qn("w:fldChar"), {qn("w:fldCharType"): "begin"}))

    run2 = paragraph.add_run()
    instr = run2._r.makeelement(qn("w:instrText"), {})
    instr.text = instruction
    run2._r.append(instr)

    run3 = paragraph.add_run()
    run3._r.append(run3._r.makeelement(
        qn("w:fldChar"), {qn("w:fldCharType"): "end"}))
    return run3


def bottom_border(target, color, size="8"):
    """
    Add a bottom rule to a paragraph or to a paragraph style.

    A Paragraph exposes its XML as ``._p``; a ParagraphStyle as ``.element``.
    """
    # Explicit None check: lxml elements are falsy when they have no children,
    # so `a or b` would silently pick the wrong one for an empty paragraph.
    element = getattr(target, "_p", None)
    if element is None:
        element = target.element
    p_pr = element.get_or_add_pPr()
    borders = p_pr.makeelement(qn("w:pBdr"), {})
    borders.append(borders.makeelement(
        qn("w:bottom"),
        {qn("w:val"): "single", qn("w:sz"): size,
         qn("w:space"): "4", qn("w:color"): color},
    ))
    p_pr.append(borders)


# ---------------------------------------------------------------------------
# Document scaffolding
# ---------------------------------------------------------------------------
def configure_page(doc):
    """A4 portrait with the standard margins."""
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH_A4
    section.page_height = PAGE_HEIGHT_A4
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT
    return section


def configure_styles(doc):
    """
    Restyle the built-in Word styles.

    Setting named styles rather than formatting individual runs is what makes
    the document navigable: Heading 1-3 drive the TOC and Word's navigation
    pane, which ad-hoc run formatting cannot do.
    """
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = GRAY_DARK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    h1 = styles["Heading 1"]
    h1.font.name = BODY_FONT
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = PRIMARY
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    bottom_border(h1, HEADER_FILL)

    h2 = styles["Heading 2"]
    h2.font.name = BODY_FONT
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = PRIMARY
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = BODY_FONT
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = GRAY_DARK
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = BODY_FONT
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = GRAY_MEDIUM
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(10)

    return doc


def add_cover_page(doc, *, title, subtitle, meta,
                   approvals: Optional[Sequence] = None,
                   logo_path: Optional[str] = None,
                   headline: Optional[Sequence] = None):
    """
    Cover page: optional logo, title, accent rule, metadata table, optional
    headline figure and approval block.

    ``meta``      - sequence of (label, value) pairs
    ``approvals`` - sequence of (heading, value) pairs, or None to omit
    ``headline``  - (value, caption, RGBColor) rendered large, or None
    """
    configure_page(doc)

    if logo_path:
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_para.add_run().add_picture(logo_path, width=Cm(4.5))
    else:
        doc.add_paragraph()

    for _ in range(2):
        doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run(title)
    run.font.name = BODY_FONT
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = PRIMARY

    para_sub = doc.add_paragraph()
    run = para_sub.add_run(subtitle)
    run.font.name = BODY_FONT
    run.font.size = Pt(13)
    run.font.color.rgb = GRAY_MEDIUM
    para_sub.paragraph_format.space_after = Pt(10)
    bottom_border(para_sub, ACCENT_FILL, size="12")

    doc.add_paragraph()

    table = doc.add_table(rows=len(meta), cols=2)
    for index, (label, value) in enumerate(meta):
        left, right = table.rows[index].cells
        left.text = str(label)
        left.width = Cm(4.5)
        style_cell(left, bold=True, size=Pt(10), color=GRAY_MEDIUM)
        right.text = "" if value is None else str(value)
        style_cell(right, size=Pt(10), color=BLACK)
    remove_table_borders(table)

    if headline:
        value, caption, color = headline
        doc.add_paragraph()
        score_para = doc.add_paragraph()
        run = score_para.add_run(str(value))
        run.font.name = BODY_FONT
        run.font.size = Pt(40)
        run.font.bold = True
        run.font.color.rgb = color

        caption_para = doc.add_paragraph()
        run = caption_para.add_run(caption)
        run.font.name = BODY_FONT
        run.font.size = Pt(12)
        run.font.color.rgb = GRAY_MEDIUM

    if approvals:
        doc.add_paragraph()
        doc.add_paragraph()
        approval = doc.add_table(rows=2, cols=len(approvals))
        approval.style = "Table Grid"
        approval.alignment = WD_TABLE_ALIGNMENT.LEFT
        for column, (heading, value) in enumerate(approvals):
            head_cell = approval.rows[0].cells[column]
            head_cell.text = str(heading)
            style_cell(head_cell, bold=True, size=Pt(9), color=GRAY_MEDIUM)
            set_cell_shading(head_cell, META_FILL)

            value_cell = approval.rows[1].cells[column]
            value_cell.text = "" if value is None else str(value)
            style_cell(value_cell, size=Pt(10), color=BLACK)

    doc.add_page_break()


def add_toc(doc, *, page_break=True):
    heading = doc.add_paragraph("Table of Contents", style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)

    para = doc.add_paragraph()
    add_field(para, r' TOC \o "1-3" \h \z \u ')

    note = doc.add_paragraph()
    run = note.add_run(
        "If the contents list appears empty, click it and press F9 to update "
        "the field (a Word behaviour for generated documents)."
    )
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = GRAY_MEDIUM

    if page_break:
        doc.add_page_break()


def configure_headers_footers(doc, *, project, doc_number, doc_date,
                              classification="Internal Use"):
    for index, section in enumerate(doc.sections):
        if index == 0:
            section.different_first_page_header_footer = True

        header = section.header
        header.is_linked_to_previous = False
        para = header.paragraphs[0]
        para.text = ""
        run = para.add_run(f"{doc_number}    |    {project}")
        run.font.name = BODY_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_MEDIUM

        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0]
        para.text = ""
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run(f"{classification}    |    Page ")
        run.font.name = BODY_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_MEDIUM

        page_run = add_field(para, " PAGE ")
        page_run.font.name = BODY_FONT
        page_run.font.size = Pt(8)
        page_run.font.color.rgb = GRAY_MEDIUM

        run = para.add_run(f"    |    {doc_date}")
        run.font.name = BODY_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY_MEDIUM


# ---------------------------------------------------------------------------
# Content helpers
# ---------------------------------------------------------------------------
def add_body(doc, text, *, bold_lead=None, space_after=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    if bold_lead:
        run = para.add_run(bold_lead)
        run.font.bold = True
        run.font.color.rgb = BLACK
    para.add_run(text)
    return para


def add_bullets(doc, items: Iterable, *, style="List Bullet"):
    for item in items:
        para = doc.add_paragraph(style=style)
        para.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            lead, rest = item
            run = para.add_run(lead)
            run.font.bold = True
            run.font.color.rgb = BLACK
            para.add_run(rest)
        else:
            para.add_run(str(item))


def add_note(doc, label, text, *, fill=NOTE_FILL):
    """A single-cell shaded callout."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]

    first = cell.paragraphs[0]
    first.text = ""
    run = first.add_run(str(label).upper())
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    first.paragraph_format.space_after = Pt(2)

    body = cell.add_paragraph()
    run = body.add_run(str(text))
    run.font.name = BODY_FONT
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY_DARK

    set_cell_shading(cell, fill)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, *, caption=None, widths=None,
              align=None, highlight=None, row_fill=None):
    """
    Render a table with the blue header row and alternating shading.

    ``align``     - per-column WD_ALIGN_PARAGRAPH, or None for left
    ``highlight`` - predicate(row) -> True to tint the row as critical
    ``row_fill``  - callable(row) -> fill hex or None, for per-row status tints
                    (takes precedence over ``highlight``)
    """
    if caption:
        para = doc.add_paragraph(caption, style="Caption")
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for column, title in enumerate(headers):
        cell = table.rows[0].cells[column]
        cell.text = str(title)
        style_cell(cell, bold=True, size=Pt(9), color=WHITE,
                   align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_shading(cell, HEADER_FILL)

    for row_index, row in enumerate(rows):
        explicit = row_fill(row) if row_fill else None
        tinted = highlight(row) if highlight else False
        for column, value in enumerate(row):
            cell = table.rows[row_index + 1].cells[column]
            cell.text = "" if value is None else str(value)
            style_cell(
                cell, size=Pt(9),
                align=(align[column] if align else None),
            )
            if explicit:
                set_cell_shading(cell, explicit)
            elif tinted:
                set_cell_shading(cell, CRITICAL_FILL)
            elif row_index % 2 == 1:
                set_cell_shading(cell, ALT_ROW_FILL)

    if widths:
        for row in table.rows:
            for column, width in enumerate(widths):
                row.cells[column].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc, image_stream, *, caption=None, width=CONTENT_WIDTH):
    """Embed a chart image, centred, with an optional numbered caption."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run().add_picture(image_stream, width=width)

    if caption:
        cap = doc.add_paragraph(caption, style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return para
