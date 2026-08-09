"""
Build the User Guide as a Word document.

Run from the repository root:

    python tools/build_user_guide_docx.py [-o OUTPUT.docx]

The scoring tables and the worked example are generated from
``src/analysis/health_score.py`` and from the sample schedule, not transcribed.
Re-run this after changing the weighting and the document follows automatically.
"""

import argparse
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from src.analysis import health_score  # noqa: E402

# ---------------------------------------------------------------------------
# Design system
# ---------------------------------------------------------------------------
COLORS = {
    "primary": RGBColor(0x00, 0x4B, 0x87),
    "white": RGBColor(0xFF, 0xFF, 0xFF),
    "black": RGBColor(0x21, 0x21, 0x21),
    "gray_dark": RGBColor(0x42, 0x42, 0x42),
    "gray_medium": RGBColor(0x75, 0x75, 0x75),
}

HEADER_FILL = "004B87"
ALT_ROW_FILL = "F5F5F5"
CRITICAL_FILL = "FADBD8"
NOTE_FILL = "EEF3F8"

PAGE_WIDTH_A4 = Cm(21.0)
PAGE_HEIGHT_A4 = Cm(29.7)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(2.0)

BODY_FONT = "Arial"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def _style_cell(cell, bold=False, size=Pt(9), color=None, align=None):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        if align is not None:
            paragraph.alignment = align
        for run in paragraph.runs:
            run.font.name = BODY_FONT
            run.font.size = size
            run.font.bold = bold
            run.font.color.rgb = color or COLORS["gray_dark"]


def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.makeelement(
        qn("w:shd"), {qn("w:fill"): hex_color, qn("w:val"): "clear"}
    )
    tc_pr.append(shading)


def _remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(borders.makeelement(
            qn(f"w:{name}"),
            {qn("w:val"): "none", qn("w:sz"): "0",
             qn("w:space"): "0", qn("w:color"): "auto"},
        ))
    tbl_pr.append(borders)


def _add_field(paragraph, instruction):
    """Insert a Word field code (used for TOC and page numbers)."""
    run = paragraph.add_run()
    run._r.append(run._r.makeelement(
        qn("w:fldChar"), {qn("w:fldCharType"): "begin"}))

    run2 = paragraph.add_run()
    instr = run2._r.makeelement(qn("w:instrText"), {})
    instr.text = instruction
    run2._r.append(instr)

    run3 = paragraph.add_run()
    run3._r.append(run3._r.makeelement(
        qn("w:fldChar"), {qn("w:fldCharType"): "end"}))
    return run3


def _bottom_border(target, color, size="8"):
    """
    Add a bottom rule to a paragraph or to a paragraph style.

    A Paragraph exposes its XML as ``._p``; a ParagraphStyle as ``.element``.
    """
    element = getattr(target, "_p", None) or target.element
    p_pr = element.get_or_add_pPr()
    borders = p_pr.makeelement(qn("w:pBdr"), {})
    borders.append(borders.makeelement(
        qn("w:bottom"),
        {qn("w:val"): "single", qn("w:sz"): size,
         qn("w:space"): "4", qn("w:color"): color},
    ))
    p_pr.append(borders)


# ---------------------------------------------------------------------------
# Document scaffolding
# ---------------------------------------------------------------------------
def configure_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = COLORS["gray_dark"]
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    h1 = styles["Heading 1"]
    h1.font.name = BODY_FONT
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = COLORS["primary"]
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    _bottom_border(h1, "004B87")

    h2 = styles["Heading 2"]
    h2.font.name = BODY_FONT
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = COLORS["primary"]
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    h3.font.name = BODY_FONT
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = COLORS["gray_dark"]
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = BODY_FONT
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = COLORS["gray_medium"]
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(10)

    return doc


def add_cover_page(doc, *, title, subtitle, project, doc_number, revision,
                   doc_date, prepared_by, checked_by, approved_by):
    section = doc.sections[0]
    section.page_width = PAGE_WIDTH_A4
    section.page_height = PAGE_HEIGHT_A4
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    for _ in range(3):
        doc.add_paragraph()

    para = doc.add_paragraph()
    run = para.add_run(title)
    run.font.name = BODY_FONT
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]

    para_sub = doc.add_paragraph()
    run = para_sub.add_run(subtitle)
    run.font.name = BODY_FONT
    run.font.size = Pt(13)
    run.font.color.rgb = COLORS["gray_medium"]
    para_sub.paragraph_format.space_after = Pt(10)
    _bottom_border(para_sub, "FFD100", size="12")

    doc.add_paragraph()

    meta = [
        ("Document No.", doc_number),
        ("Revision", revision),
        ("Applies to", project),
        ("Date", doc_date),
        ("Audience", "Planners, schedulers, project controls"),
        ("Classification", "Internal Use"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    for index, (label, value) in enumerate(meta):
        left, right = table.rows[index].cells
        left.text = label
        left.width = Cm(4.5)
        _style_cell(left, bold=True, size=Pt(10), color=COLORS["gray_medium"])
        right.text = value
        _style_cell(right, size=Pt(10), color=COLORS["black"])
    _remove_table_borders(table)

    doc.add_paragraph()
    doc.add_paragraph()

    approval = doc.add_table(rows=2, cols=3)
    approval.style = "Table Grid"
    approval.alignment = WD_TABLE_ALIGNMENT.LEFT
    for column, (heading, value) in enumerate(
            zip(("Prepared by", "Checked by", "Approved by"),
                (prepared_by, checked_by, approved_by))):
        head_cell = approval.rows[0].cells[column]
        head_cell.text = heading
        _style_cell(head_cell, bold=True, size=Pt(9),
                    color=COLORS["gray_medium"])
        _set_cell_shading(head_cell, "D6E4F0")

        value_cell = approval.rows[1].cells[column]
        value_cell.text = value
        _style_cell(value_cell, size=Pt(10), color=COLORS["black"])

    doc.add_page_break()


def add_toc(doc):
    heading = doc.add_paragraph("Table of Contents", style="Heading 1")
    heading.paragraph_format.space_before = Pt(0)

    para = doc.add_paragraph()
    _add_field(para, r' TOC \o "1-3" \h \z \u ')

    note = doc.add_paragraph()
    run = note.add_run(
        "If the contents list appears empty, click it and press F9 to update "
        "the field (a Word behaviour for generated documents)."
    )
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = COLORS["gray_medium"]

    doc.add_page_break()


def configure_headers_footers(doc, *, project, doc_number, doc_date):
    for index, section in enumerate(doc.sections):
        if index == 0:
            section.different_first_page_header_footer = True

        header = section.header
        header.is_linked_to_previous = False
        para = header.paragraphs[0]
        para.text = ""
        run = para.add_run(f"{doc_number}    |    {project}")
        run.font.name = BODY_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS["gray_medium"]

        footer = section.footer
        footer.is_linked_to_previous = False
        para = footer.paragraphs[0]
        para.text = ""
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run("Internal Use    |    Page ")
        run.font.name = BODY_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS["gray_medium"]

        page_run = _add_field(para, " PAGE ")
        page_run.font.name = BODY_FONT
        page_run.font.size = Pt(8)
        page_run.font.color.rgb = COLORS["gray_medium"]

        run = para.add_run(f"    |    {doc_date}")
        run.font.name = BODY_FONT
        run.font.size = Pt(8)
        run.font.color.rgb = COLORS["gray_medium"]


# ---------------------------------------------------------------------------
# Content helpers
# ---------------------------------------------------------------------------
def add_body(doc, text, *, bold_lead=None, space_after=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    if bold_lead:
        run = para.add_run(bold_lead)
        run.font.bold = True
        run.font.color.rgb = COLORS["black"]
    para.add_run(text)
    return para


def add_bullets(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            lead, rest = item
            run = para.add_run(lead)
            run.font.bold = True
            run.font.color.rgb = COLORS["black"]
            para.add_run(rest)
        else:
            para.add_run(item)


def add_note(doc, label, text):
    """A single-cell shaded callout."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]

    first = cell.paragraphs[0]
    first.text = ""
    run = first.add_run(label.upper())
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    first.paragraph_format.space_after = Pt(2)

    body = cell.add_paragraph()
    run = body.add_run(text)
    run.font.name = BODY_FONT
    run.font.size = Pt(9.5)
    run.font.color.rgb = COLORS["gray_dark"]

    _set_cell_shading(cell, NOTE_FILL)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, *, caption=None, widths=None,
              align=None, highlight=None):
    """
    Render a table with the blue header row and alternating shading.

    ``align``     - per-column WD_ALIGN_PARAGRAPH, or None for left
    ``highlight`` - predicate(row) -> True to tint the row
    """
    if caption:
        para = doc.add_paragraph(caption, style="Caption")
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for column, title in enumerate(headers):
        cell = table.rows[0].cells[column]
        cell.text = str(title)
        _style_cell(cell, bold=True, size=Pt(9), color=COLORS["white"],
                    align=WD_ALIGN_PARAGRAPH.LEFT)
        _set_cell_shading(cell, HEADER_FILL)

    for row_index, row in enumerate(rows):
        tinted = highlight(row) if highlight else False
        for column, value in enumerate(row):
            cell = table.rows[row_index + 1].cells[column]
            cell.text = "" if value is None else str(value)
            _style_cell(
                cell, size=Pt(9),
                align=(align[column] if align else None),
            )
            if tinted:
                _set_cell_shading(cell, CRITICAL_FILL)
            elif row_index % 2 == 1:
                _set_cell_shading(cell, ALT_ROW_FILL)

    if widths:
        for row in table.rows:
            for column, width in enumerate(widths):
                row.cells[column].width = width

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


# ---------------------------------------------------------------------------
# Live data
# ---------------------------------------------------------------------------
def component_rows():
    """The weighting table, straight from the scoring module."""
    rows = []
    for component in health_score.COMPONENTS:
        zero_at = (f"{component.zero_at:g}%" if not component.higher_is_better
                   else f"≤{component.zero_at:g}")
        rows.append([
            component.dcma or "-",
            component.label,
            component.weight,
            component.describe_target(),
            zero_at,
            "Yes" if component.critical else "",
        ])
    return rows


def worked_example():
    """Score a real schedule so the example can never be out of date."""
    from src.analysis.dcma_analyzer import DCMAAnalyzer
    from src.analysis.metrics_calculator import MetricsCalculator
    from src.parsers.schedule_parser import ScheduleParser

    source = ROOT / "data" / "Schedule export.csv"
    if not source.exists():
        source = ROOT / "data" / "sample_schedule.csv"

    parsed = ScheduleParser().parse_csv(source.read_bytes(), source.name)
    if not parsed.get("success"):
        return None, None, None

    results = DCMAAnalyzer(parsed).analyze()
    metrics = MetricsCalculator(parsed, results["metrics"]).calculate_all_metrics()
    health = metrics["health_score"]

    rows = []
    for component in health["components"]:
        if component["score"] is None:
            measured, score, contribution = "n/a", "n/a", "-"
        else:
            measured = f"{component['value']}{component['unit']}"
            score = f"{component['score']:.0f}"
            contribution = f"{component['score'] * component['weight'] / 100:.2f}"
        rows.append([
            component["label"], measured, component["target"],
            score, component["weight"], contribution,
        ])

    return parsed["total_activities"], health, rows


def recovery_ranking(health):
    """Points recoverable per check, largest first."""
    scored = [
        (c["label"], (100 - c["score"]) * c["weight"] / 100)
        for c in health["components"]
        if c["score"] is not None and c["score"] < 100
    ]
    return sorted(scored, key=lambda item: item[1], reverse=True)


# ---------------------------------------------------------------------------
# Document body
# ---------------------------------------------------------------------------
RIGHT = WD_ALIGN_PARAGRAPH.RIGHT
LEFT = WD_ALIGN_PARAGRAPH.LEFT
CENTER = WD_ALIGN_PARAGRAPH.CENTER


def build_part_one(doc):
    doc.add_heading("Part 1 — Using the Application", level=1)

    doc.add_heading("What it does", level=2)
    add_body(doc,
             "The application reads a Primavera P6 schedule exported as CSV and "
             "assesses it against the DCMA 14-Point Schedule Assessment, "
             "returning metrics, issues, prioritised recommendations and a "
             "single 0–100 health score.")
    add_body(doc,
             "It does not assess the GAO Schedule Assessment Guide, and it does "
             "not read .xer files — CSV export only.")

    doc.add_heading("Roles and access", level=2)
    add_table(
        doc,
        ["Role", "Can do"],
        [["Admin", "Upload, analyse, delete, generate reports, manage users"],
         ["Viewer", "Read-only: dashboards, comparisons, report downloads"]],
        widths=[Cm(3.5), Cm(12.5)],
    )
    add_body(doc,
             "Sessions expire after 60 minutes of inactivity. There are no "
             "shared or demo accounts — each person gets their own, created by "
             "an admin under Settings → User Management.")

    doc.add_heading("Preparing the P6 export", level=2)
    add_body(doc, "Export as CSV. The file is rejected without the required "
                  "columns.")
    add_table(
        doc,
        ["Column", "Status", "What depends on it"],
        [["Activity ID", "Required", "Identity of every activity"],
         ["Activity Name", "Required", "Reporting and issue lists"],
         ["Activity Status", "Required", "BEI, status breakdown"],
         ["Start / Finish", "Required", "Durations, date checks, project span"],
         ["Total Float", "Required", "Float analysis, negative float, CPLI"],
         ["Duration Type", "Required", "—"],
         ["Predecessor Details", "Recommended",
          "Every logic metric: leads, lags, relationship types, missing logic"],
         ["Successor Details", "Recommended", "Missing successors, open ends"],
         ["WBS Code", "Recommended", "WBS analysis and per-area breakdowns"],
         ["At Completion Duration", "Recommended", "Duration checks"],
         ["Free Float", "Recommended", "Float analysis"],
         ["Primary Constraint", "Recommended", "Constraint checks"],
         ["Activity Type", "Recommended", "Milestone handling"],
         ["Resource Names", "Recommended", "Resource-loading check"]],
        caption="Table 1 — Columns the parser looks for",
        widths=[Cm(4.6), Cm(2.9), Cm(9.0)],
        highlight=lambda row: row[0] == "Predecessor Details",
    )

    add_note(doc, "The one column to get right",
             "Export 'Predecessor Details', not just 'Predecessors'. The plain "
             "column carries activity IDs only — no relationship type, no lag. "
             "If that is all you export, the application warns you and assumes "
             "every relationship is Finish-to-Start with zero lag. Leads and "
             "lags then read as zero no matter what the schedule actually "
             "contains.")

    add_note(doc, "Dates",
             "P6 writes dates in the exporting machine's locale with nothing to "
             "say which one it used, and 03/04/2025 is genuinely ambiguous. The "
             "parser inspects every date in the file and picks the only "
             "interpretation consistent with the data, then reports what it "
             "chose. If a file is entirely ambiguous it says so rather than "
             "guessing silently. Export as YYYY-MM-DD and the question "
             "disappears.")

    doc.add_heading("The workflow", level=2)

    steps = [
        ("Step 1 — Sign in", None,
         ["Repeated failed attempts lock the account temporarily."]),
        ("Step 2 — Create or select a project",
         "Upload Schedule → 1. Select or Create Project",
         ["A project groups successive revisions of the same schedule, and "
          "project codes are unique. Create it once, then upload each revision "
          "against it — that is what makes the Comparison page useful."]),
        ("Step 3 — Upload and analyse",
         "Upload Schedule → 2. Upload Schedule File → 3. Upload and Analyze",
         ["Preview the first ten rows to confirm the columns landed where you "
          "expect, then run the analysis. A 1,200-activity schedule takes a few "
          "seconds.",
          "Read the warnings. They are not noise: they report which date format "
          "was detected, which columns were missing, which relationships could "
          "not be parsed and which dates could not be read. A schedule that "
          "analyses successfully with five warnings may have been assessed on "
          "incomplete data."]),
        ("Step 4 — Review the analysis", "Analysis Dashboard",
         ["Seven tabs, summarised in Table 2 below. On the Overview tab, "
          "'How this score is calculated' expands to the full breakdown for "
          "that schedule: every check, its measured value, its target, its "
          "weight and its score."]),
        ("Step 5 — Compare revisions", "Comparison",
         ["Select two versions of the same project for metrics side by side "
          "with the movement between them — improvements green, regressions "
          "red. Needs at least two schedules."]),
        ("Step 6 — Generate reports", "Reports",
         ["DOCX — executive summary: cover page, health score, DCMA checklist, "
          "CPLI and BEI, issues, recommendations, methodology appendix. For "
          "stakeholders.",
          "Excel — full detail: summary, issues, complete activity list, logic "
          "breakdown, recommendations. For working the problem. Exports are "
          "recorded in the audit log."]),
    ]

    for title, where, paragraphs in steps:
        doc.add_heading(title, level=3)
        if where:
            para = doc.add_paragraph()
            run = para.add_run(where)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = COLORS["gray_medium"]
            para.paragraph_format.space_after = Pt(4)
        for text in paragraphs:
            add_body(doc, text)

    doc.add_heading("Reading the dashboard", level=2)
    add_table(
        doc,
        ["Tab", "What it holds"],
        [["Overview", "Health score, CPLI, BEI, headline counts, data-quality warnings"],
         ["Detailed Metrics", "Every DCMA check with its count, percentage and status"],
         ["Float Analysis", "Float distribution, negative float, float by WBS"],
         ["WBS Analysis", "Per-area breakdown and comparison"],
         ["Issues", "Findings by severity with the affected activities"],
         ["Recommendations", "Prioritised actions with impact and effort"],
         ["Activities", "The parsed activity table"]],
        caption="Table 2 — Analysis Dashboard tabs",
        widths=[Cm(4.2), Cm(12.3)],
    )


def build_part_two(doc):
    doc.add_page_break()
    doc.add_heading("Part 2 — How the Health Score Is Set", level=1)

    add_body(doc,
             "The health score is a weighted average of twelve DCMA checks, "
             "expressed 0–100:")

    formula = doc.add_paragraph()
    formula.alignment = CENTER
    run = formula.add_run(
        "Score  =  Σ (check score × weight)  ÷  Σ (weight of checks with data)")
    run.font.name = BODY_FONT
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    formula.paragraph_format.space_before = Pt(8)
    formula.paragraph_format.space_after = Pt(12)

    add_note(doc, "Where the numbers come from",
             "The thresholds are DCMA's. The weights are not — DCMA defines "
             "pass/fail criteria for each point, not a composite score. The "
             "weights below are this application's assessment of relative "
             "severity. They are published here so they can be challenged and "
             "changed deliberately rather than buried in code.")

    doc.add_heading("How one check scores", level=2)
    add_body(doc,
             "Each check is measured as a percentage of the schedule, then "
             "scored 0–100: full marks at or better than the DCMA target, zero "
             "at or beyond a defined bound, declining linearly in between.")
    add_table(
        doc,
        ["Measured value", "Check score"],
        [["3%", "100"], ["5% (the target)", "100"], ["10%", "80"],
         ["17.5%", "50"], ["30% (the zero bound)", "0"], ["45%", "0"]],
        caption="Table 3 — Logic completeness, target ≤5%, zero at 30%",
        widths=[Cm(8.0), Cm(8.5)],
        align=[LEFT, RIGHT],
    )

    doc.add_heading("The twelve checks", level=2)
    add_table(
        doc,
        ["DCMA", "Check", "Weight", "Target", "Scores 0 at", "Critical"],
        component_rows() + [["", "Total", sum(
            c.weight for c in health_score.COMPONENTS), "", "", ""]],
        caption="Table 4 — Weighting (src/analysis/health_score.py)",
        widths=[Cm(1.6), Cm(5.6), Cm(2.0), Cm(2.2), Cm(2.6), Cm(2.0)],
        align=[CENTER, LEFT, RIGHT, RIGHT, RIGHT, CENTER],
        highlight=lambda row: row[5] == "Yes",
    )

    add_body(doc,
             "because every other network metric depends on it: float, the "
             "critical path, CPLI and the relationship checks are all "
             "meaningless in a schedule that is not properly linked.",
             bold_lead="Logic completeness carries the largest weight ")
    add_body(doc,
             "in any valid network — the project's own beginning and end — and "
             "are excluded from the missing-logic measurement. Without that "
             "exclusion a clean three-activity chain would measure 67% missing "
             "logic.",
             bold_lead="One open start and one open finish are expected ")

    doc.add_heading("Checks with no data", level=2)
    add_body(doc,
             "A check whose input column is absent is marked n/a and excluded, "
             "and the remaining weights are renormalised over what could "
             "actually be measured. It does not score as a silent pass: a "
             "schedule exported without resource names is not credited with "
             "being fully resourced.")

    doc.add_heading("Two rules beyond the average", level=2)
    add_body(doc,
             "A plain weighted average has a weakness: a catastrophic failure "
             "on one check costs only that check's weight, so a schedule with "
             "half its relationships as leads would still average out near "
             "Excellent. Two rules correct that.")

    doc.add_heading("Data-sufficiency gates", level=3)
    add_body(doc, "Applied when the schedule cannot be meaningfully assessed "
                  "at all.")
    add_table(
        doc,
        ["Condition", "Score capped at"],
        [["No relationship data whatsoever",
          f"{health_score.GATES['no_relationships'][0]:.0f}"],
         ["More than 50% of activities missing logic",
          f"{health_score.GATES['logic_unusable'][0]:.0f}"]],
        widths=[Cm(11.0), Cm(5.5)],
        align=[LEFT, RIGHT],
    )

    doc.add_heading("Critical-check ceilings", level=3)
    add_body(doc,
             "Each critical check — marked Yes in Table 4 — that fails "
             "outright, meaning it scores below "
             f"{health_score.CRITICAL_FAIL_BELOW:.0f}, costs one rating band.")
    ceiling_rows = []
    for minimum, ceiling, _ in health_score.CEILINGS:
        rating = next(name for threshold, name, _ in health_score.RATING_BANDS
                      if ceiling >= threshold)
        label = f"{minimum}" if minimum < len(health_score.CEILINGS) \
            else f"{minimum} or more"
        ceiling_rows.append([label, f"{ceiling:.0f}", rating])
    add_table(
        doc,
        ["Critical checks failed", "Capped at", "Best possible rating"],
        ceiling_rows,
        caption="Table 5 — Rating ceilings",
        widths=[Cm(5.5), Cm(4.0), Cm(7.0)],
        align=[CENTER, RIGHT, LEFT],
    )
    add_body(doc,
             "Ceilings only ever lower a score. If the weighted average already "
             "sits below the ceiling, nothing changes and nothing is reported. "
             "Every cap that is applied appears on the dashboard with its "
             "reason — a reduced score is never left unexplained.")

    doc.add_heading("Rating bands", level=2)
    band_rows = []
    bands = health_score.RATING_BANDS
    for index, (threshold, name, _) in enumerate(bands):
        upper = 100 if index == 0 else bands[index - 1][0] - 1
        band_rows.append([f"{threshold}–{upper}", name])
    add_table(
        doc,
        ["Score", "Rating"],
        band_rows,
        widths=[Cm(6.0), Cm(10.5)],
        align=[CENTER, LEFT],
    )

    # ---- Worked example ------------------------------------------------
    doc.add_heading("Worked example", level=2)
    activities, health, rows = worked_example()
    if not health:
        add_body(doc, "Sample schedule unavailable at build time.")
        return

    add_body(doc,
             f"A real {activities:,}-activity P6 export, every figure as the "
             f"application produced it.")

    add_table(
        doc,
        ["Check", "Measured", "Target", "Score", "Weight", "Contribution"],
        rows + [["Total", "", "", "", health["applicable_weight"],
                 f"{health['score']:.1f}"]],
        caption=f"Table 6 — Worked example ({activities:,} activities)",
        widths=[Cm(5.2), Cm(2.4), Cm(2.2), Cm(1.9), Cm(2.2), Cm(2.6)],
        align=[LEFT, RIGHT, RIGHT, RIGHT, RIGHT, RIGHT],
    )

    result = doc.add_paragraph()
    run = result.add_run(f"Score {health['score']:.1f} — {health['rating']}")
    run.font.name = BODY_FONT
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = COLORS["primary"]
    result.paragraph_format.space_after = Pt(10)

    if health["caps"]:
        for cap in health["caps"]:
            add_body(doc, cap, bold_lead="Cap applied: ")
    else:
        add_body(doc, "No cap was applied to this schedule.")

    ranking = recovery_ranking(health)
    total_lost = sum(points for _, points in ranking)
    add_body(doc,
             ", ".join(f"{label} {points:.1f}" for label, points in ranking[:5])
             + f". These account for {total_lost:.1f} of the "
             f"{100 - health['score']:.1f} points lost, which is the practical "
             f"value of the breakdown: it ranks the remedial work by how much "
             f"it is actually worth.",
             bold_lead="Points recoverable, largest first: ")

    doc.add_heading("What the score is, and is not", level=2)
    add_body(doc, "It is:", space_after=2)
    add_bullets(doc, [
        "A transparent, reproducible summary of twelve DCMA checks",
        "A way to track one schedule across revisions",
        "A means of directing attention to the worst areas first",
    ])
    add_body(doc, "It is not:", space_after=2)
    add_bullets(doc, [
        "A DCMA-defined figure — DCMA publishes no composite score",
        "A substitute for reading the Issues tab",
        "Comparable against scores produced by any other tool",
        "Meaningful where parse warnings flagged missing relationship or date data",
    ])
    add_body(doc,
             "Use it to answer 'is this revision better than the last, and "
             "where should I look first'. Do not use it as an acceptance "
             "criterion on its own.")

    doc.add_heading("Changing the weighting", level=2)
    add_note(doc, "Before changing any weight",
             "The weights are a business judgement, not a constant. They live "
             "in src/analysis/health_score.py in the COMPONENTS table, with the "
             "gates and ceilings directly below it. Changing them changes every "
             "score the application reports, including on analyses already "
             "issued to clients. Re-run 'python -m pytest "
             "tests/test_health_score.py' afterwards — it asserts that the "
             "weights sum to 100, that scoring stays proportional, and that a "
             "schedule with no logic never outranks a well-linked one.")


def build_troubleshooting(doc):
    doc.add_page_break()
    doc.add_heading("Troubleshooting", level=1)
    add_table(
        doc,
        ["Message", "Cause and fix"],
        [["Missing required columns",
          "Re-export from P6 including the required columns in Table 1"],
         ["Date format is ambiguous",
          "Every date could read either way. Re-export as YYYY-MM-DD, or set "
          "APP_DATE_ORDER to day or month"],
         ["Dates are inconsistent",
          "The file mixes day-first and month-first values. Re-export as "
          "YYYY-MM-DD"],
         ["n values could not be read as a date",
          "Those dates were left empty. Check the source cells for text or a "
          "stray format"],
         ["Using 'Predecessors' column",
          "Only activity IDs were available. Re-export with Predecessor "
          "Details for accurate lead and lag metrics"],
         ["No relationship data",
          "The export has no predecessor or successor columns. Logic metrics "
          "are unavailable and the score is capped at 25"],
         ["Duplicate column headers",
          "Two columns share a name; re-export with unique headers"],
         ["File exceeds the limit",
          "Default ceiling is 50 MB — see DEPLOYMENT.md"],
         ["Analysis failed, with a reference code",
          "Quote the reference to your administrator; the full detail is in "
          "the application log"]],
        caption="Table 7 — Common messages",
        widths=[Cm(5.4), Cm(11.1)],
    )

    doc.add_heading("Further reading", level=2)
    add_bullets(doc, [
        ("README.md", " — installation and first run"),
        ("DEPLOYMENT.md", " — production configuration, backups, security posture"),
        ("USER_GUIDE.md", " — the source this document is generated from"),
        ("src/analysis/health_score.py", " — the scoring weights, gates and ceilings"),
    ])


def build(output_path: Path) -> Path:
    doc = Document()
    configure_styles(doc)

    today = date.today().strftime("%d/%m/%Y")
    add_cover_page(
        doc,
        title="Schedule Quality Analyzer",
        subtitle="User Guide and Health Score Methodology",
        project="Schedule Quality Analyzer",
        doc_number="SQA-GUI-0001",
        revision="Rev. 01",
        doc_date=today,
        prepared_by="Project Controls",
        checked_by="Planning Manager",
        approved_by="Project Manager",
    )
    add_toc(doc)
    build_part_one(doc)
    build_part_two(doc)
    build_troubleshooting(doc)
    configure_headers_footers(
        doc,
        project="Schedule Quality Analyzer",
        doc_number="SQA-GUI-0001",
        doc_date=today,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "-o", "--output",
        default=str(ROOT / "Schedule_Quality_Analyzer_User_Guide.docx"),
        help="Output .docx path",
    )
    args = parser.parse_args()

    written = build(Path(args.output))
    size_kb = written.stat().st_size / 1024
    print(f"Wrote {written}  ({size_kb:.0f} KB)")


if __name__ == "__main__":
    main()
